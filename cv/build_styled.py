"""Build an ATS-safe DOCX that matches the visual style of the original CV.

Reproduces the original look (Cambria 11pt, centered name, ruled section
headings, right-aligned dates) WITHOUT layout tables -- dates are placed with a
right tab stop so the document stays single-column for parsers.

Usage:
    python build_styled.py resume_source.yaml Dimosthenis-Gkontolias-Resume.docx
"""

import sys

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BODY_FONT = "Cambria"
BODY_SIZE = 10.5
NAME_SIZE = 15
HEADING_SIZE = 10.5
CONTACT_SIZE = 9.5
META_SIZE = 9.5  # company / location / role -- secondary to the bold title

# Minimum visual gap (points) between an entry's left text and its right-aligned
# date. Below this the date visually collides with the title.
MIN_GAP_PT = 24.0

PAGE_W = 8.5
MARGIN_LR = 0.5
MARGIN_TB = 0.4
TEXT_W = PAGE_W - 2 * MARGIN_LR  # right tab stop position


def set_border(paragraph, size, space=1):
    """Add a bottom border to a paragraph (size is in eighths of a point)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def tight(paragraph, before=0, after=0, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return paragraph


def run(paragraph, text, bold=False, italic=False, size=BODY_SIZE):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = BODY_FONT
    return r


def rich(paragraph, text, size=BODY_SIZE):
    """Render text where **...** marks an inline bold span.

    Bold is an attention budget, not decoration: Lever strips it entirely, so
    it must never be the only thing carrying meaning.
    """
    for i, chunk in enumerate(text.split("**")):
        if chunk:
            run(paragraph, chunk, bold=(i % 2 == 1), size=size)


def heading(doc, text):
    p = doc.add_paragraph()
    tight(p, before=10, after=4)
    run(p, text.upper(), bold=True, size=HEADING_SIZE)
    set_border(p, 6)
    return p


def measure(parts):
    """Width in points of (text, bold, size) parts, using real Cambria metrics."""
    try:
        from PIL import ImageFont
    except ImportError:
        return 0.0
    total = 0.0
    for text, bold, size in parts:
        path = "C:/Windows/Fonts/cambriab.ttf" if bold else "C:/Windows/Fonts/cambria.ttc"
        total += ImageFont.truetype(path, int(size * 10)).getlength(text) / 10.0
    return total


def entry_line(doc, left_parts, right_text):
    """One line: title left, date pushed right by a tab stop.

    left_parts is a list of (text, bold) -- bold parts render at BODY_SIZE,
    the rest at META_SIZE so the title stays dominant and the line stays short.
    """
    p = doc.add_paragraph()
    tight(p, before=6, after=1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(TEXT_W), WD_TAB_ALIGNMENT.RIGHT)
    for text, bold in left_parts:
        run(p, text, bold=bold, size=BODY_SIZE if bold else META_SIZE)
    if right_text:
        run(p, "\t" + right_text, bold=True)

    measured = [(t, b, BODY_SIZE if b else META_SIZE) for t, b in left_parts]
    gap = TEXT_W * 72 - measure(measured) - measure([(right_text, True, BODY_SIZE)])
    if gap < MIN_GAP_PT:
        print(
            f"  WARNING: only {gap:.0f}pt before the date on "
            f"{left_parts[0][0][:40]!r} -- shorten it (want >={MIN_GAP_PT:.0f}pt)"
        )
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    tight(p, before=1, after=1, line=1.03)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.22)
    pf.first_line_indent = Inches(-0.14)
    rich(p, "•  " + text)
    return p


def build(src_path, out_path):
    data = yaml.safe_load(open(src_path, encoding="utf-8"))
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    s = doc.sections[0]
    s.page_width, s.page_height = Inches(PAGE_W), Inches(11)
    s.left_margin = s.right_margin = Inches(MARGIN_LR)
    s.top_margin = s.bottom_margin = Inches(MARGIN_TB)

    # --- header -----------------------------------------------------------
    p = doc.add_paragraph()
    tight(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, data["name"], bold=True, size=NAME_SIZE)

    if data.get("headline"):
        p = doc.add_paragraph()
        tight(p, after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, data["headline"], italic=True, size=CONTACT_SIZE)

    contact = [data.get("location"), data.get("email"), data.get("phone")]
    p = doc.add_paragraph()
    tight(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "  |  ".join(x for x in contact if x), size=CONTACT_SIZE)

    links = [l if isinstance(l, str) else l.get("url", "") for l in data.get("links", [])]
    p = doc.add_paragraph()
    tight(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "  |  ".join(links), size=CONTACT_SIZE)
    set_border(p, 12)

    # --- summary ----------------------------------------------------------
    if data.get("summary"):
        heading(doc, "Summary")
        for line in data["summary"]:
            p = doc.add_paragraph()
            tight(p, after=2, line=1.03)
            rich(p, line)

    # --- skills -----------------------------------------------------------
    if data.get("skills"):
        heading(doc, "Technical Skills")
        for category, values in data["skills"].items():
            # Category labels stay regular weight: bolding six of them would
            # compete with the six real eye-tracking fixation points.
            p = doc.add_paragraph()
            tight(p, after=2, line=1.03)
            run(p, f"{category}: {', '.join(values)}")

    # --- experience -------------------------------------------------------
    if data.get("experience"):
        heading(doc, "Work Experience")
        for e in data["experience"]:
            left = [(e["title"], True)]
            tail = " @ " + e["company"]
            if e.get("location"):
                tail += f", {e['location']}"
            left.append((tail, False))
            dates = f"{e.get('start', '')} - {e.get('end', '')}".strip(" -")
            entry_line(doc, left, dates)
            for b in e.get("bullets", []):
                bullet(doc, b)

    # --- projects ---------------------------------------------------------
    if data.get("projects"):
        heading(doc, "Projects")
        if data.get("projects_note"):
            p = doc.add_paragraph()
            tight(p, after=1)
            run(p, data["projects_note"], italic=True, size=META_SIZE)
        for pr in data["projects"]:
            left = [(pr["name"], True)]
            if pr.get("role"):
                left.append((f", {pr['role']}", False))
            entry_line(doc, left, pr.get("dates", ""))
            for b in pr.get("bullets", []):
                bullet(doc, b)

    # --- education --------------------------------------------------------
    if data.get("education"):
        heading(doc, "Education")
        for ed in data["education"]:
            left = [(ed["degree"], True), (" - " + ed["institution"], False)]
            if ed.get("location"):
                left.append((f", {ed['location']}", False))
            entry_line(doc, left, ed.get("dates", ""))
            for d in ed.get("details", []):
                bullet(doc, d)

    # --- custom sections --------------------------------------------------
    for sec in data.get("custom_sections", []):
        heading(doc, sec["heading"])
        for item in sec.get("items", []):
            if item.get("meta"):
                entry_line(doc, [(item["title"], False)], item["meta"])
            else:
                p = doc.add_paragraph()
                tight(p, after=0)
                run(p, item["title"])
            for b in item.get("bullets", []):
                bullet(doc, b)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
